from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
for p in [str(PROJECT_ROOT), str(PROJECT_ROOT / "backend")]:
    if p not in sys.path:
        sys.path.insert(0, p)

from app.parsing.common import clean_text
from app.parsing.excel_parser import parse_xlsx
from app.parsing.metadata import detect_period, detect_unit, find_clause_no, parse_filename
from app.parsing.models import ParsedDocument
from app.parsing.table_evidence import compact_header, iter_table_evidence
from app.parsing.word_parser import parse_docx


def make_document(path: Path, doc_id: str = "test_doc") -> ParsedDocument:
    return ParsedDocument(
        doc_id=doc_id,
        source_seq=1,
        file_name=path.name,
        file_type=path.suffix.lstrip("."),
        file_size=path.stat().st_size,
        sha256="0" * 64,
        local_path=path.name,
        absolute_path=path,
        title=path.stem,
    )


class MetadataTest(unittest.TestCase):
    def test_filename_and_financial_fields(self) -> None:
        info = parse_filename(Path("001_监管页面标题_附件表.xlsx"))
        self.assertEqual(info["doc_id"], "nfra_att_001")
        self.assertEqual(info["source_page_title"], "监管页面标题")
        self.assertEqual(info["attachment_title"], "附件表")
        self.assertEqual(find_clause_no("第二十条 商业银行应当建立制度"), "第二十条")
        self.assertEqual(detect_period("2025年三季度"), "2025Q3")
        self.assertEqual(detect_unit("单位：亿元"), "亿元")
        self.assertEqual(clean_text(4806061.691218953), "4806061.691218953")


class WordParserTest(unittest.TestCase):
    def test_docx_keeps_heading_clause_and_table(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "001_监管办法_监管办法.docx"
            doc = Document()
            doc.add_heading("监管办法", level=1)
            doc.add_paragraph("第一条 商业银行应当建立管理制度。")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "指标"
            table.cell(0, 1).text = "2025年"
            table.cell(1, 0).text = "资本充足率"
            table.cell(1, 1).text = "12.5%"
            doc.save(path)

            bundle = parse_docx(path, make_document(path))
            self.assertTrue(any(block.block_type == "heading" for block in bundle.blocks))
            self.assertTrue(any(block.clause_no == "第一条" for block in bundle.blocks))
            self.assertEqual(len(bundle.tables), 1)
            cells = list(bundle.tables[0].iter_cells())
            self.assertTrue(any(cell.cell_ref == "R2C2" and cell.display_value == "12.5%" for cell in cells))


class ExcelParserTest(unittest.TestCase):
    def test_xlsx_multi_header_formula_and_location(self) -> None:
        import openpyxl

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "002_统计表_统计表.xlsx"
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "资本指标"
            sheet.merge_cells("A1:C1")
            sheet["A1"] = "2025年三季度资本监管指标"
            sheet["C2"] = "单位：%"
            sheet.append(["指标", "2025年二季度", "2025年三季度"])
            sheet.append(["资本充足率", 12.1, 12.5])
            sheet["C5"] = "=C4-B4"
            workbook.save(path)
            workbook.close()

            bundle = parse_xlsx(path, make_document(path))
            self.assertEqual(len(bundle.tables), 1)
            cells = list(bundle.tables[0].iter_cells())
            value = next(cell for cell in cells if cell.cell_ref == "C4")
            formula = next(cell for cell in cells if cell.cell_ref == "C5")
            self.assertEqual(value.normalized_value, 12.5)
            self.assertEqual(value.period, "2025Q3")
            self.assertEqual(formula.formula, "=C4-B4")


class TableEvidenceTest(unittest.TestCase):
    def test_groups_cells_into_retrieval_row_and_keeps_locations(self) -> None:
        rows = [
            {
                "doc_id": "doc1",
                "title": "地区保费表",
                "table_id": "table1",
                "sheet_name": "数据",
                "table_name": "地区保费表",
                "row_index": 4,
                "col_index": 2,
                "cell_ref": "B4",
                "value": "北京",
                "row_header": "",
                "col_header": "地区保费表 / 地区",
                "is_header": 0,
            },
            {
                "doc_id": "doc1",
                "title": "地区保费表",
                "table_id": "table1",
                "sheet_name": "数据",
                "table_name": "地区保费表",
                "row_index": 4,
                "col_index": 3,
                "cell_ref": "C4",
                "value": "100.25",
                "normalized_value": "100.25",
                "row_header": "北京",
                "col_header": "地区保费表 / 合计",
                "period": "2025-09",
                "unit": "亿元",
                "is_header": 0,
            },
        ]
        evidence = list(iter_table_evidence(rows, max_cells_per_evidence=20))
        self.assertEqual(len(evidence), 1)
        self.assertEqual(evidence[0]["cell_range"], "B4:C4")
        self.assertEqual(evidence[0]["metric_name"], "北京")
        self.assertIn("合计=100.25", evidence[0]["retrieval_text"])
        self.assertEqual(compact_header("地区保费表 / 单位：亿元 / 合计", title="地区保费表"), "合计")


if __name__ == "__main__":
    unittest.main()
